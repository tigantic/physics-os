"""
Document Extractor
==================

Extracts data from Word documents (DOC, DOCX) and text files.
Parses specification documents, equipment schedules, and project notes.
"""

import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
import io

try:
    from . import (
        BaseExtractor, ExtractionResult, ExtractedField,
        ExtractionConfidence
    )
    from ..units import detect_unit_system, UnitSystem
except ImportError:
    import sys
    sys.path.insert(0, str(Path(__file__).parent))
    sys.path.insert(0, str(Path(__file__).parent.parent))
    from extractors import (
        BaseExtractor, ExtractionResult, ExtractedField,
        ExtractionConfidence
    )
    from units import detect_unit_system, UnitSystem


class DocumentExtractor(BaseExtractor):
    """
    Extract HVAC data from text documents.
    
    Handles:
    - DOCX files (python-docx)
    - Plain text files
    - Markdown files
    - Rich text format
    """
    
    SUPPORTED_EXTENSIONS = [".docx", ".doc", ".txt", ".md", ".rtf"]
    SUPPORTED_MIME_TYPES = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
        "text/plain",
        "text/markdown",
    ]
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check for required dependencies."""
        self._has_docx = False
        
        try:
            import docx
            self._has_docx = True
        except ImportError:
            pass
    
    def extract(self, file_path) -> ExtractionResult:
        """Extract data from a document file."""
        # Handle both string and Path objects
        if isinstance(file_path, str):
            file_path = Path(file_path)
        
        if not file_path.exists():
            return self._create_empty_result(
                file_path.name, "document",
                f"File not found: {file_path}"
            )
        
        with open(file_path, "rb") as f:
            data = f.read()
        
        return self.extract_from_bytes(data, file_path.name)
    
    def extract_from_bytes(self, data: bytes, filename: str) -> ExtractionResult:
        """Extract data from document bytes."""
        file_hash = self._compute_hash(data)
        fields: Dict[str, ExtractedField] = {}
        warnings: List[str] = []
        errors: List[str] = []
        raw_text = ""
        
        ext = Path(filename).suffix.lower()
        
        try:
            if ext in [".docx"]:
                if not self._has_docx:
                    return self._create_empty_result(
                        filename, "document",
                        "python-docx not installed. Run: pip install python-docx"
                    )
                
                raw_text = self._extract_docx(data)
                
            elif ext == ".doc":
                # Legacy .doc requires different library
                warnings.append("Legacy .doc format - limited extraction")
                try:
                    raw_text = data.decode("utf-8", errors="ignore")
                except UnicodeDecodeError:
                    raw_text = ""
                    
            elif ext in [".txt", ".md"]:
                raw_text = data.decode("utf-8", errors="ignore")
            
            else:
                raw_text = data.decode("utf-8", errors="ignore")
            
            # Detect unit system
            unit_system, unit_confidence = detect_unit_system(raw_text)
            
            # Extract structured fields
            fields.update(self._extract_dimensions(raw_text))
            fields.update(self._extract_latex_si(raw_text))  # LaTeX/SI units
            fields.update(self._extract_hvac_data(raw_text))
            fields.update(self._extract_project_info(raw_text))
            fields.update(self._extract_spec_data(raw_text))
            fields.update(self._extract_table_data(raw_text))
            
            return ExtractionResult(
                success=True,
                file_name=filename,
                file_type="document",
                file_hash=file_hash,
                extracted_at=datetime.now(),
                detected_unit_system=unit_system,
                unit_confidence=unit_confidence,
                fields=fields,
                warnings=warnings,
                errors=errors,
                raw_text=raw_text,
            )
            
        except Exception as e:
            return self._create_empty_result(
                filename, "document",
                f"Document extraction failed: {str(e)}"
            )
    
    def _extract_docx(self, data: bytes) -> str:
        """Extract text from DOCX file."""
        import docx
        
        doc = docx.Document(io.BytesIO(data))
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            text_parts.append("[TABLE]")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
            text_parts.append("[/TABLE]")
        
        return "\n".join(text_parts)
    
    def _extract_spec_data(self, text: str) -> Dict[str, ExtractedField]:
        """Extract specification-specific data."""
        fields = {}
        
        # Design conditions
        design_patterns = {
            "summer_outdoor_db": [
                r"summer\s*(?:outdoor|outside)\s*(?:db|dry\s*bulb)[:\s]+(\d+(?:\.\d+)?)\s*°?([FC])?",
                r"cooling\s*design[:\s]+(\d+(?:\.\d+)?)\s*°?([FC])?",
            ],
            "winter_outdoor_db": [
                r"winter\s*(?:outdoor|outside)\s*(?:db|dry\s*bulb)[:\s]+(\d+(?:\.\d+)?)\s*°?([FC])?",
                r"heating\s*design[:\s]+(\d+(?:\.\d+)?)\s*°?([FC])?",
            ],
            "indoor_cooling_setpoint": [
                r"(?:cooling|summer)\s*(?:indoor|inside|room)\s*(?:temp|setpoint)[:\s]+(\d+(?:\.\d+)?)\s*°?([FC])?",
                r"(?:indoor|inside|room)\s*(?:cooling|summer)[:\s]+(\d+(?:\.\d+)?)\s*°?([FC])?",
            ],
            "indoor_heating_setpoint": [
                r"(?:heating|winter)\s*(?:indoor|inside|room)\s*(?:temp|setpoint)[:\s]+(\d+(?:\.\d+)?)\s*°?([FC])?",
                r"(?:indoor|inside|room)\s*(?:heating|winter)[:\s]+(\d+(?:\.\d+)?)\s*°?([FC])?",
            ],
            "relative_humidity": [
                r"(?:relative\s*)?humidity[:\s]+(\d+(?:\.\d+)?)\s*%?",
                r"(\d+(?:\.\d+)?)\s*%\s*RH",
            ],
        }
        
        for field_name, patterns in design_patterns.items():
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    value = float(match.group(1))
                    unit = match.group(2) if match.lastindex >= 2 else None
                    
                    fields[field_name] = ExtractedField(
                        name=field_name,
                        value=value,
                        confidence=ExtractionConfidence.MEDIUM,
                        original_text=match.group(0),
                        unit=unit,
                    )
                    break
        
        # Ventilation requirements
        vent_patterns = [
            r"ventilation[:\s]+(\d+(?:\.\d+)?)\s*(?:CFM|cfm)",
            r"outside\s*air[:\s]+(\d+(?:\.\d+)?)\s*(?:CFM|cfm)",
            r"OA[:\s]+(\d+(?:\.\d+)?)\s*(?:CFM|cfm)",
        ]
        
        for pattern in vent_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                fields["ventilation_cfm"] = ExtractedField(
                    name="ventilation_cfm",
                    value=float(match.group(1)),
                    confidence=ExtractionConfidence.MEDIUM,
                    original_text=match.group(0),
                    unit="CFM",
                )
                break
        
        # Equipment model numbers
        model_patterns = [
            r"model[:\s#]+([A-Z0-9\-]+)",
            r"(?:AHU|RTU|FCU)[:\s#]*([A-Z0-9\-]+)",
        ]
        
        models = []
        for pattern in model_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            models.extend(matches)
        
        if models:
            fields["equipment_models"] = ExtractedField(
                name="equipment_models",
                value=list(set(models)),
                confidence=ExtractionConfidence.MEDIUM,
            )
        
        return fields
    
    def _extract_table_data(self, text: str) -> Dict[str, ExtractedField]:
        """Extract data from tables in the document."""
        fields = {}
        
        # Find table sections
        table_pattern = r"\[TABLE\](.*?)\[/TABLE\]"
        tables = re.findall(table_pattern, text, re.DOTALL)
        
        for i, table in enumerate(tables):
            rows = table.strip().split("\n")
            if len(rows) >= 2:
                # Check if it's a schedule table
                header = rows[0].lower()
                
                if any(term in header for term in ["cfm", "airflow", "supply", "return"]):
                    # Likely a diffuser/equipment schedule
                    schedule_data = []
                    for row in rows[1:]:
                        cells = [c.strip() for c in row.split("|")]
                        if len(cells) >= 2:
                            schedule_data.append(cells)
                    
                    if schedule_data:
                        fields[f"schedule_table_{i}"] = ExtractedField(
                            name=f"schedule_table_{i}",
                            value=schedule_data,
                            confidence=ExtractionConfidence.LOW,
                        )
        
        return fields
